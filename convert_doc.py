#!/usr/bin/env python3
"""
Convert markdown deliverables to .docx or .xlsx
Usage: python3 convert_doc.py <input.md> <output_path>
Output format is determined by content analysis:
  - If the file is primarily tables → .xlsx
  - Otherwise → .docx
"""
import sys, os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

KPMG_BLUE = RGBColor(0x00, 0x33, 0x8D)
KPMG_NAVY = RGBColor(0x1B, 0x1F, 0x3B)
KPMG_DARK = RGBColor(0x33, 0x33, 0x33)
KPMG_GRAY = RGBColor(0x6D, 0x6E, 0x71)

def parse_md_tables(text):
    """Extract markdown tables as list of (header_row, data_rows)."""
    tables = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:\-]+$', lines[i+1].strip()):
            headers = [c.strip() for c in line.strip('|').split('|')]
            rows = []
            i += 2  # skip separator
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            tables.append((headers, rows))
        else:
            i += 1
    return tables

def is_spreadsheet_content(text):
    """Determine if content is primarily tabular."""
    tables = parse_md_tables(text)
    total_table_rows = sum(len(rows) for _, rows in tables)
    total_lines = len([l for l in text.split('\n') if l.strip()])
    return total_table_rows > 20 and total_table_rows > total_lines * 0.4

def md_to_docx(md_text, output_path, title="Document"):
    doc = Document()
    
    # Style setup
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = KPMG_DARK
    style.paragraph_format.space_after = Pt(6)
    
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = KPMG_BLUE if level <= 2 else KPMG_NAVY
        hs.font.bold = True
        hs.font.size = Pt([0, 18, 14, 12, 11][level])
    
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Headings
        if stripped.startswith('####'):
            doc.add_heading(stripped.lstrip('#').strip(), level=4)
            i += 1; continue
        if stripped.startswith('###'):
            doc.add_heading(stripped.lstrip('#').strip(), level=3)
            i += 1; continue
        if stripped.startswith('##'):
            doc.add_heading(stripped.lstrip('#').strip(), level=2)
            i += 1; continue
        if stripped.startswith('#'):
            doc.add_heading(stripped.lstrip('#').strip(), level=1)
            i += 1; continue
        
        # Horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a thin line
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1; continue
        
        # Tables
        if '|' in stripped and i + 1 < len(lines) and re.match(r'^[\s|:\-]+$', lines[i+1].strip()):
            headers = [c.strip() for c in stripped.strip('|').split('|')]
            rows = []
            i += 2
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(row)
                i += 1
            
            ncols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j] if j < ncols else None
                if cell:
                    cell.text = clean_md_inline(h)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
            
            for ri, row in enumerate(rows):
                for j, val in enumerate(row):
                    if j < ncols:
                        cell = table.rows[ri + 1].cells[j]
                        cell.text = clean_md_inline(val)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(10)
            
            doc.add_paragraph()  # spacing after table
            continue
        
        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = clean_md_inline(stripped[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1; continue
        
        # Numbered lists
        m = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m:
            text = clean_md_inline(m.group(2))
            p = doc.add_paragraph(text, style='List Number')
            i += 1; continue
        
        # Empty lines
        if not stripped:
            i += 1; continue
        
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1
    
    doc.save(output_path)

def clean_md_inline(text):
    """Remove markdown inline formatting for plain text contexts."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text

def add_formatted_runs(paragraph, text):
    """Add runs with bold/italic formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = KPMG_BLUE
        else:
            paragraph.add_run(part)

def md_to_xlsx(md_text, output_path, title="Data"):
    wb = openpyxl.Workbook()
    
    header_fill = PatternFill(start_color='00338D', end_color='00338D', fill_type='solid')
    header_font = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
    cell_font = Font(name='Calibri', size=10, color='333333')
    thin_border = Border(
        left=Side(style='thin', color='E5E7EB'),
        right=Side(style='thin', color='E5E7EB'),
        top=Side(style='thin', color='E5E7EB'),
        bottom=Side(style='thin', color='E5E7EB')
    )
    
    tables = parse_md_tables(md_text)
    
    if not tables:
        ws = wb.active
        ws.title = title[:31]
        ws['A1'] = md_text
        wb.save(output_path)
        return
    
    # Remove default sheet
    wb.remove(wb.active)
    
    for idx, (headers, rows) in enumerate(tables):
        # Try to get a sheet name from the heading before the table
        sheet_name = f"Table {idx+1}"
        ws = wb.create_sheet(title=sheet_name[:31])
        
        # Write headers
        for j, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=j, value=clean_md_inline(h))
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='left', vertical='center')
            cell.border = thin_border
        
        # Write data
        for ri, row in enumerate(rows, 2):
            for j, val in enumerate(row, 1):
                cell = ws.cell(row=ri, column=j, value=clean_md_inline(val))
                cell.font = cell_font
                cell.border = thin_border
                # Try to convert numbers
                clean_val = val.strip().replace(',', '').replace('$', '').replace('%', '')
                try:
                    cell.value = float(clean_val) if '.' in clean_val else int(clean_val)
                    cell.number_format = '#,##0.00' if '.' in clean_val else '#,##0'
                    if '$' in val:
                        cell.number_format = '$#,##0' if '.' not in clean_val else '$#,##0.00'
                    if '%' in val:
                        cell.number_format = '0.0%'
                        cell.value = cell.value / 100 if abs(cell.value) < 100 else cell.value
                except (ValueError, TypeError):
                    cell.value = clean_md_inline(val)
        
        # Auto-fit columns
        for j in range(1, len(headers) + 1):
            max_len = max(len(str(ws.cell(row=r, column=j).value or '')) for r in range(1, len(rows) + 2))
            ws.column_dimensions[openpyxl.utils.get_column_letter(j)].width = min(max_len + 4, 50)
    
    wb.save(output_path)

def convert(input_path, output_dir):
    with open(input_path, 'r') as f:
        content = f.read()
    
    basename = os.path.splitext(os.path.basename(input_path))[0]
    
    if is_spreadsheet_content(content):
        out_path = os.path.join(output_dir, basename + '.xlsx')
        md_to_xlsx(content, out_path, basename)
        return out_path, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    else:
        out_path = os.path.join(output_dir, basename + '.docx')
        md_to_docx(content, out_path, basename)
        return out_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: convert_doc.py <input.md> <output_dir>")
        sys.exit(1)
    out_path, mime = convert(sys.argv[1], sys.argv[2])
    print(f"{out_path}|{mime}")
